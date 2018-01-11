from docx import Document
from Utils import *
from .Processor import Processor

from win32com.client import Dispatch


class DocProcessor(Processor):
    """
    只能处理docx，doc类型需要用win32com转换
    空格作分隔符来分隔段落连续字符串
    表格遍历所有单元格
    """

    def __init__(self, file, process_impl):
        super(DocProcessor, self).__init__(file, process_impl)
        self._name = "word processor"

    def doc_to_docx(self):
        # 调用win32com的方法打开word文件并另存为docx格式，
        # 用win32com打开是word程序独占的操作，打开到关闭期间不能手动用word打开文档，也无法多进程，希望有更好的解决方案
        try:
            word = Dispatch("Word.Application")
            word.visible = 0
            word.displayalerts = 0
            doc = word.Documents.Open(FileName=self.file)
            filePathOut = change_ext_to(self.file, "docx")
            doc.SaveAs(filePathOut, 16)
        except Exception as e:
            raise e
        finally:
            try:
                doc.Close()
                word.Quit()
            except Exception as e:
                raise e

        return filePathOut

    def read_data(self):
        doc = Document(self.file)
        self.gather_paragraphs(doc)
        self.gather_tables(doc)

    def gather_tables(self, doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if has_mobile_num(cell.text):
                        self._src_data.append(cell.text)

                    self.gather_tables(cell)

    def gather_paragraphs(self, doc):
        for p in doc.paragraphs:
            items = p.text.split()
            self._src_data.extend(items)
